#!/usr/bin/env python3
"""
Runtime dual-write: Notion row + Drive docx with dedupe_key idempotency.
Use from CLI or invoked by /nizam-governor-push skill after local POP commit.
"""
from __future__ import annotations

import argparse
import io
import json
import sys
from datetime import datetime
from typing import Any

import requests

from nizam_governor_lib import (
    build_drive_record_filename,
    build_drive_record_path,
    build_drive_service,
    build_meeting_path,
    build_weekly_review_path,
    check_privacy_gate,
    compute_dedupe_key,
    date_from_captured_at,
    emit_receipt,
    ensure_drive_path_chain,
    list_drive_children,
    load_config,
    normalize_percent,
    payload_hash,
    runtime_receipt,
    slugify,
    stage_human_only_fields,
    utc_now_iso,
    get_notion_token,
)

NOTION_VERSION = "2022-06-28"


def normalize_record(raw: dict[str, Any], config: dict[str, Any]) -> dict[str, Any]:
    """Build governor_runtime_record from CLI/skill payload."""
    session_map = config["session_type_to_record_type"]
    lane_defaults = config["default_lane_by_session_type"]

    lane = raw.get("lane")
    record_type = raw.get("type")
    slug = raw.get("slug")
    captured_at = raw.get("captured_at") or utc_now_iso()

    if raw.get("session_type") and not record_type:
        st = raw["session_type"]
        record_type = session_map.get(st, "CheckIn")
        lane = lane or lane_defaults.get(st, "Recovery")

    lane = lane or "Recovery"
    record_type = record_type or "CheckIn"
    slug = slug or slugify(raw.get("notion_title") or raw.get("source_artifact") or "record")
    date_str = date_from_captured_at(captured_at)
    dedupe_key = raw.get("dedupe_key") or compute_dedupe_key(lane, record_type, date_str, slug)

    notion_primary = raw.get("notion_primary", "witness")
    routing = config.get("artifact_routing", [])
    for rule in routing:
        if raw.get("session_type") == rule.get("session_type"):
            notion_primary = rule.get("notion_primary", notion_primary)
            lane = rule.get("lane", lane)
            record_type = rule.get("type", record_type)
            break

    return {
        "dedupe_key": dedupe_key,
        "lane": lane,
        "type": record_type,
        "slug": slug,
        "captured_at": captured_at,
        "notion_title": raw.get("notion_title")
        or f"{lane} — {record_type} ({date_str})",
        "notion_primary": notion_primary,
        "notion_secondary": raw.get("notion_secondary"),
        "notion_payload": raw.get("notion_payload") or {},
        "drive_narrative": raw.get("drive_narrative") or "",
        "drive_path_override": raw.get("drive_path_override"),
        "repo_commit": raw.get("repo_commit"),
        "source_artifact": raw.get("source_artifact"),
        "human_only_staged": raw.get("human_only_staged") or [],
        "operator_confirmed_externalize": raw.get("operator_confirmed_externalize", False),
        "privacy_classification": raw.get("privacy_classification", "strict_local"),
    }


def resolve_drive_path(record: dict[str, Any]) -> str:
    if record.get("drive_path_override"):
        return record["drive_path_override"]
    lane = record["lane"]
    rtype = record["type"]
    date_str = date_from_captured_at(record["captured_at"])
    slug = record["slug"]
    if rtype == "WeeklyReview":
        dt = datetime.strptime(date_str, "%Y-%m-%d")
        return build_weekly_review_path(dt.isocalendar().year, dt.isocalendar().week)
    if rtype == "Meeting":
        return build_meeting_path(date_str, slug)
    filename = build_drive_record_filename(lane, rtype, date_str, slug)
    return build_drive_record_path(lane, filename)


def build_docx_bytes(record: dict[str, Any], notion_page_id: str | None) -> bytes:
    from docx import Document

    doc = Document()
    doc.add_heading(record.get("notion_title") or record["dedupe_key"], level=1)
    doc.add_paragraph(f"notion_page_id: {notion_page_id or ''}")
    doc.add_paragraph(f"dedupe_key: {record['dedupe_key']}")
    doc.add_paragraph(f"captured_at: {record['captured_at']}")
    doc.add_paragraph(f"repo_commit: {record.get('repo_commit') or ''}")
    doc.add_paragraph("")
    doc.add_paragraph(record["drive_narrative"])
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def notion_headers(token: str) -> dict[str, str]:
    return {
        "Authorization": f"Bearer {token}",
        "Notion-Version": NOTION_VERSION,
        "Content-Type": "application/json",
    }


def notion_query_by_dedupe(
    token: str, database_id: str, dedupe_key: str, prop_name: str
) -> str | None:
    """Query database; return page_id if dedupe_key exists."""
    url = f"https://api.notion.com/v1/databases/{database_id}/query"
    body = {
        "filter": {
            "property": prop_name,
            "rich_text": {"equals": dedupe_key},
        }
    }
    resp = requests.post(url, headers=notion_headers(token), json=body, timeout=30)
    if resp.status_code == 400:
        # Property may be title type
        body = {"filter": {"property": prop_name, "title": {"equals": dedupe_key}}}
        resp = requests.post(url, headers=notion_headers(token), json=body, timeout=30)
    if not resp.ok:
        return None
    results = resp.json().get("results", [])
    return results[0]["id"] if results else None


def notion_create_page(
    token: str,
    database_id: str,
    title: str,
    properties: dict[str, Any],
) -> str:
    props = {
        "Name": {"title": [{"text": {"content": title[:2000]}}]},
    }
    props.update(properties)
    body = {"parent": {"database_id": database_id}, "properties": props}
    resp = requests.post(
        "https://api.notion.com/v1/pages",
        headers=notion_headers(token),
        json=body,
        timeout=30,
    )
    resp.raise_for_status()
    return resp.json()["id"]


def notion_update_page(token: str, page_id: str, properties: dict[str, Any]) -> None:
    resp = requests.patch(
        f"https://api.notion.com/v1/pages/{page_id}",
        headers=notion_headers(token),
        json={"properties": properties},
        timeout=30,
    )
    resp.raise_for_status()


def notion_build_properties(
    record: dict[str, Any], config: dict[str, Any], drive_url: str | None = None
) -> dict[str, Any]:
    req = config["notion"]["required_properties"]
    props: dict[str, Any] = {}
    dk = record["dedupe_key"]
    props[req["dedupe_key"]] = {"rich_text": [{"text": {"content": dk}}]}
    props[req["captured_at"]] = {
        "date": {"start": record["captured_at"], "time_zone": "UTC"}
    }
    if drive_url:
        props[req["drive_link"]] = {"url": drive_url}
    payload = record.get("notion_payload") or {}
    for k, v in payload.items():
        if k in config.get("human_only_fields", []):
            continue
        if isinstance(v, (int, float)) and 0 <= v <= 100:
            v = normalize_percent(v)
        if isinstance(v, float) and v <= 1:
            props[k] = {"number": v}
        elif isinstance(v, (int, float)):
            props[k] = {"number": v}
        elif isinstance(v, str):
            props[k] = {"rich_text": [{"text": {"content": v[:2000]}}]}
    return props


def notion_audit_log(
    token: str,
    config: dict[str, Any],
    *,
    event_type: str,
    entity_id: str,
    outcome: str,
    phash: str,
) -> bool:
    ds = config["notion"]["data_sources"]["audit_log"]["data_source_id"]
    try:
        notion_create_page(
            token,
            ds,
            f"{event_type} {utc_now_iso()}",
            {
                "event_type": {"rich_text": [{"text": {"content": event_type}}]},
                "time": {"date": {"start": utc_now_iso()}},
                "entity_id": {"rich_text": [{"text": {"content": entity_id}}]},
                "payload_hash": {"rich_text": [{"text": {"content": phash}}]},
                "outcome": {"rich_text": [{"text": {"content": outcome}}]},
            },
        )
        return True
    except requests.RequestException:
        return False


def upload_docx(service, parent_id: str, filename: str, content: bytes, file_id: str | None) -> str:
    from googleapiclient.http import MediaIoBaseUpload

    media = MediaIoBaseUpload(
        io.BytesIO(content),
        mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        resumable=False,
    )
    if file_id:
        service.files().update(fileId=file_id, media_body=media, supportsAllDrives=True).execute()
        return file_id
    meta = {"name": filename, "parents": [parent_id]}
    return (
        service.files()
        .create(body=meta, media_body=media, fields="id, webViewLink", supportsAllDrives=True)
        .execute()
    )


def find_drive_file_by_name(service, parent_id: str, name: str) -> dict[str, str] | None:
    children = list_drive_children(service, parent_id)
    if name in children:
        fid = children[name]["id"]
        meta = service.files().get(fileId=fid, fields="webViewLink", supportsAllDrives=True).execute()
        return {"id": fid, "url": meta.get("webViewLink")}
    return None


def dual_write(payload: dict[str, Any], *, dry_run: bool = False) -> dict[str, Any]:
    config = load_config()
    human_only = config.get("human_only_fields", [])
    _, staged_top = stage_human_only_fields(dict(payload), human_only)
    record = normalize_record(payload, config)
    cleaned_payload, staged = stage_human_only_fields(
        record.get("notion_payload") or {}, human_only
    )
    record["notion_payload"] = cleaned_payload
    record["human_only_staged"] = list(
        set(record.get("human_only_staged", []) + staged + staged_top)
    )

    allowed, reason = check_privacy_gate(
        record.get("privacy_classification"),
        record.get("operator_confirmed_externalize", False),
    )
    if not allowed:
        return runtime_receipt(
            "FAILED",
            "CREATE",
            record["dedupe_key"],
            failed_stage="himayah_gate",
            notes=reason,
            human_only_staged=record["human_only_staged"],
        )

    notion_token = get_notion_token()
    drive_service = build_drive_service()
    ds_key = record["notion_primary"]
    ds_id = config["notion"]["data_sources"][ds_key]["data_source_id"]
    ds_name = config["notion"]["data_sources"][ds_key]["name"]
    prop_dedupe = config["notion"]["required_properties"]["dedupe_key"]
    drive_path = resolve_drive_path(record)
    root_id = config["targets"]["drive_root_id"]

    mode = "CREATE"
    page_id: str | None = None
    drive_url: str | None = None
    drive_file_id: str | None = None

    if dry_run:
        return runtime_receipt(
            "OK",
            mode,
            record["dedupe_key"],
            notion_db=ds_name,
            notion_ds=ds_id,
            notes=f"Dry-run: would write {drive_path}",
            human_only_staged=record["human_only_staged"],
        )

    # Stage 2: CHECK Notion
    if notion_token:
        page_id = notion_query_by_dedupe(notion_token, ds_id, record["dedupe_key"], prop_dedupe)
        if page_id:
            mode = "UPDATE"
    else:
        return runtime_receipt(
            "FAILED",
            "CREATE",
            record["dedupe_key"],
            failed_stage="notion_check",
            notes="NOTION_TOKEN not set",
        )

    # Stage 3: CHECK Drive
    if not drive_service:
        return runtime_receipt(
            "FAILED",
            mode,
            record["dedupe_key"],
            notion_db=ds_name,
            notion_ds=ds_id,
            notion_page_id=page_id,
            failed_stage="drive_check",
            notes="GOOGLE_APPLICATION_CREDENTIALS not set",
        )

    parent_id, filename = ensure_drive_path_chain(drive_service, root_id, drive_path)
    existing_drive = find_drive_file_by_name(drive_service, parent_id, filename)
    if existing_drive:
        drive_file_id = existing_drive["id"]
        mode = "UPDATE"

    # Stage 4: WRITE Notion
    props = notion_build_properties(record, config)
    try:
        if page_id:
            notion_update_page(notion_token, page_id, props)
        else:
            page_id = notion_create_page(
                notion_token,
                ds_id,
                record["notion_title"],
                props,
            )
    except requests.RequestException as e:
        return runtime_receipt(
            "FAILED",
            mode,
            record["dedupe_key"],
            failed_stage="notion_write",
            notes=str(e),
            human_only_staged=record["human_only_staged"],
        )

    # Stage 5: WRITE Drive
    try:
        docx = build_docx_bytes(record, page_id)
        result = upload_docx(
            drive_service, parent_id, filename, docx, drive_file_id
        )
        if isinstance(result, dict):
            drive_file_id = result["id"]
            drive_url = result.get("webViewLink")
        else:
            drive_file_id = result
            meta = drive_service.files().get(
                fileId=drive_file_id, fields="webViewLink", supportsAllDrives=True
            ).execute()
            drive_url = meta.get("webViewLink")
    except Exception as e:
        return runtime_receipt(
            "FAILED",
            mode,
            record["dedupe_key"],
            notion_db=ds_name,
            notion_ds=ds_id,
            notion_page_id=page_id,
            failed_stage="drive_write",
            notes=str(e),
            human_only_staged=record["human_only_staged"],
        )

    # Stage 6: DriveLink back
    drivelink_ok = False
    try:
        link_props = notion_build_properties(record, config, drive_url=drive_url)
        notion_update_page(notion_token, page_id, link_props)
        drivelink_ok = True
    except requests.RequestException:
        drivelink_ok = False

    # Stage 7: Audit
    audit_ok = notion_audit_log(
        notion_token,
        config,
        event_type="dual_write",
        entity_id=page_id or record["dedupe_key"],
        outcome="OK" if drivelink_ok else "partial",
        phash=payload_hash(record),
    )

    return runtime_receipt(
        "OK" if drivelink_ok else "FAILED",
        mode,
        record["dedupe_key"],
        notion_db=ds_name,
        notion_ds=ds_id,
        notion_page_id=page_id,
        drive_folder=drive_path.rsplit("/", 1)[0],
        drive_filename=filename,
        drive_url=drive_url,
        drivelink_written_back=drivelink_ok,
        audit_logged=audit_ok,
        human_only_staged=record["human_only_staged"],
        failed_stage=None if drivelink_ok else "drivelink_writeback",
        notes="Dual-write complete" if drivelink_ok else "Drive written; DriveLink writeback failed",
    )


def main() -> int:
    parser = argparse.ArgumentParser(description="NIZAM runtime dual-write")
    parser.add_argument("--input", "-i", help="JSON file path (stdin if omitted)")
    parser.add_argument("--dry-run", action="store_true")
    args = parser.parse_args()

    if args.input:
        with open(args.input, encoding="utf-8") as f:
            payload = json.load(f)
    else:
        payload = json.load(sys.stdin)

    receipt = dual_write(payload, dry_run=args.dry_run)
    print(emit_receipt(receipt))
    return 0 if receipt["status"] == "OK" else 1


if __name__ == "__main__":
    sys.exit(main())
